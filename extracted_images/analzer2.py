from docx import Document
from docx2python import docx2python
import os

def analyze_docx(file_path) -> str:

    doc = Document(file_path)
    

    print("Paragraphs and Styles:")
    for para in doc.paragraphs:
        print(f"Text: {para.text}")
        print(f"Style: {para.style.name}")
        print("-" * 50)
    
    print("\nTables:")
    for table_idx, table in enumerate(doc.tables):
        print(f"Table {table_idx + 1}:")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            print("\t".join(row_data))
            return "\t".join(row_data)
        print("-" * 50)
        
    
    core_props = doc.core_properties
    print("\nMetadata:")
    print(f"Author: {core_props.author}")
    print(f"Created: {core_props.created}")
    print(f"Modified: {core_props.modified}")
    print("-" * 50)
    
    doc_content = docx2python(file_path)
    print("\nHeaders:")
    for header in doc_content.header:
        print(header)
        print("-" * 50)
    
    print("\nFooters:")
    for footer in doc_content.footer:
        print(footer)
        print("-" * 50)
    
    print("\nImages:")
    image_dir = "extracted_images"
    os.makedirs(image_dir, exist_ok=True)
    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            image_part = doc.part.rels[rel].target_part
            image_data = image_part.blob
            image_filename = os.path.basename(image_part.partname)
            image_path = os.path.join(image_dir, image_filename)
            with open(image_path, "wb") as img_file:
                img_file.write(image_data)
            print(f"Extracted image saved to: {image_path}")
    print("-" * 50)


