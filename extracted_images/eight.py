# import os
# import time
# import sys
# import asyncio
# from typing import List, Optional, Dict, Any
# import re
# from pydantic import BaseModel, Field
# from pydantic_ai import Agent
# #Groq 
# from pydantic_ai.models.groq import GroqModel
# from dotenv import load_dotenv
# import traceback
# from pynput.mouse import Button, Controller as MouseController
# import logging
# from helpers import GrammarChecker
# from models import DocumentContent, RetypedDocument
# from typer import RealKeyboardTyper


# from analyzer import analyze_docx

# load_dotenv()
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# class DocumentRetyper:
#     """Class to retype documents with real keyboard typing"""
#     def __init__(self):
#         self.document_retyper = None
#         self.keyboard_typer = RealKeyboardTyper(delay=0.01)  # Slightly slower for reliability
        
#     async def async_init(self):
#         """Async initialization method"""
#         self.document_retyper = Agent(
#             # 'google-gla:gemini-2.0-flash',
#             'groq:llama-3.3-70b-versatile',
#             result_type=RetypedDocument,
#             system_prompt=(
#                 "You are a document retyper assistant. Your task is to carefully read and retype"
#                 " documents, preserving their exact original format and content. "
#                 "BOLD AND ITALIC TEXT SHOULD BE PRESERVED. "
#                  "For formatting, use the following markers:"
#             " **bold text** for bold,"
#             " __underlined text__ for underline,"
#             " _italic text_ for italic."
#             "use the addition functions like the ""addition"" function like toggle_bold, toggle_underline,toggle_italic "
#                 "USE THE REQUIRED TOOLS TO MAKE A CLEAR PROGRESS AND FASTER ERROR FREE RETYPING."
#                 "BULLETS, NUMBERING LIKE 1.2.3 ETC, AND PARAGRAPH SPACING SHOULD BE PRESERVED."
#                 "TYPE EXACT HOW THE DOCUMENT IS NO MORE STAFF ADDED IN IT AND BE CAREFUL."
#                 "Do not analyze, summarize, or modify the document in any way. Do not skip empty lines"
#                 " or spaces. Simply reproduce the document exactly as provided, maintaining all"
#                 "formatting, HEADINGS, bullet points, paragraph breaks, and structure."
#                 "Make sure that the Numbering and Bullet Points are preserved. And also incase ` is used in the document"
#                 " Make sure to preserve all newlines and paragraph spacing exactly as in the original."
#             ),
#         )
#         return self

#     async def retype_document_with_real_typing(self, document_text: str, typing_position=None):
#         """Function to retype a document using real keyboard inputs with formatting support"""
#         content = DocumentContent(text=document_text)
        
#         # Just show a simple message in terminal
#         print(f"Reading document content...")
        
#         # Run the agent silently
#         result = await self.document_retyper.run(
#             f"Please retype the following document exactly as it is, preserving all formatting, spacing,"
#             f" line breaks, and paragraph structure: {content.text}"
#         )
        
#         # Prepare user for typing
#         print(f"Preparing to type the document with real keyboard inputs.")
#         print(f"Please open the target file/editor and place your cursor where typing should begin.")
#         print(f"You have 5 seconds to get ready...")
        
#         # Countdown
#         for i in range(5, 0, -1):
#             print(f"{i}...")
#             await asyncio.sleep(1)
        
#         print("Typing now! You can move the mouse, the system will try to maintain focus...")
        
#         # Perform real keyboard typing with formatting support
#         try:
#             # Check if the content likely has formatting markers
#             if "**" in result.data.content or "__" in result.data.content or "_" in result.data.content:
#                 # Use the enhanced typing with formatting method
#                 await self.keyboard_typer.type_with_formatting(result.data.content, typing_position)
#             else:
#                 # Use the regular typing with verification method
#                 await self.keyboard_typer.type_with_verification(result.data.content, typing_position, True)
#             print(f"\nDocument successfully retyped using real keyboard inputs!")
#         except Exception as e:
#             print(f"Error during typing: {str(e)}")
#             traceback.print_exc()
        
#         return result.data.content

#     async def display_document_info(self, file_path: str):
#         """Display basic document info without typing the content"""
#         # Check if it's a docx file

#         file_ext = analyze_docx(file_path)
#         print(f"File extension: {file_ext}")

#         if file_path.endswith('.docx'):
#             import docx
#             doc = docx.Document(file_path)
#             # Count paragraphs and characters in docx
#             paragraphs = len(doc.paragraphs)
#             char_count = sum(len(paragraph.text) for paragraph in doc.paragraphs)
#             print(f"Document contains {paragraphs} paragraphs and {char_count} characters.")
#             print("Starting retyping process...\n")
#         else:
#             # For regular text files
#             with open(file_path, 'r') as f:
#                 content = f.read()
                
#             print(f"Document loaded: {file_path}")
#             line_count = content.count('\n') + 1
#             char_count = len(content)
#             print(f"Document contains {line_count} lines and {char_count} characters.")
#             print("Starting retyping process...\n")

# # Function to extract text from docx with better formatting preservation
# def extract_text_from_docx(docx_path):
#     """Extract text from docx file with enhanced formatting preservation"""
#     import docx
#     doc = docx.Document(docx_path)
#     full_text = []
    
#     # Process paragraphs with enhanced formatting
#     for para in doc.paragraphs:
#         # Don't skip empty paragraphs - they're part of formatting
#         # Handle heading styles
#         if para.style.name.startswith('Heading'):
#             # Add appropriate formatting based on heading level
#             level = int(para.style.name.replace('Heading', '')) if para.style.name != 'Heading' else 1
#             prefix = '#' * level + ' ' if level > 0 else ''
#             full_text.append(f"{prefix}{para.text}")
#         else:
#             # Regular paragraph - include even if empty to preserve spacing
#             full_text.append(para.text)
    
#     # Process tables if any
#     for table in doc.tables:
#         table_rows = []
#         for row in table.rows:
#             row_text = []
#             for cell in row.cells:
#                 row_text.append(cell.text.strip())
#             table_rows.append(" | ".join(row_text))
#         full_text.append("\n".join(table_rows))
    
#     # Join with paragraph spacing preserved
#     return "\n\n".join(full_text)

# # Example usage
# if __name__ == "__main__":
#     import asyncio
#     import os
#     import docx
#     import random
#     from docx import Document
    
#     async def main():
#         # Create and initialize the processor
#         retyper = DocumentRetyper()
#         await retyper.async_init()
        
#         # Handle command line arguments if provided
#         if len(sys.argv) > 1:
#             file_path = sys.argv[1]
#             if not os.path.exists(file_path):
#                 print(f"Error: File {file_path} not found.")
#                 return
            
#             output_file = "retyped_output.txt"
#             if len(sys.argv) > 2:
#                 output_file = sys.argv[2]
#         else:
#             # Default sample document path
#             file_path = "doc.docx"
#             output_file = "my_data.txt"
            
#         # Show basic info about the document
#         await retyper.display_document_info(file_path)
        
#         # Get document text from file
#         if file_path.endswith('.docx'):
#             document_text = extract_text_from_docx(file_path)
#         else:
#             with open(file_path, 'r') as f:
#                 document_text = f.read()
        
#         # Prepare for typing
#         print("Please open a text editor where you want the typing to occur.")
#         input("Press Enter when you're ready to begin typing...")
        
#         # Get cursor position (optional)
#         print("Place your cursor where typing should begin and keep it there.")
#         print("Typing will start in 3 seconds...")
#         await asyncio.sleep(3)
        
#         # Current mouse position will be used as the focus point
#         mouse = MouseController()
#         cursor_position = mouse.position
        
#         # Run the retyping directly to file
#         try:
#             retyped_content = await retyper.retype_document_with_real_typing(document_text, cursor_position)
#             print(f"Document has been retyped with real keyboard inputs!")
#         except Exception as e:
#             print(f"Error occurred: {str(e)}")
#             traceback.print_exc()
    
#     asyncio.run(main())
import os
import time
import sys
import asyncio
from typing import List, Optional, Dict, Any
import re
from pydantic import BaseModel, Field
from pydantic_ai import Agent
#Groq 
from pydantic_ai.models.groq import GroqModel
from dotenv import load_dotenv
import traceback
from pydantic.error_wrappers import ValidationError
from pynput.mouse import Button, Controller as MouseController
from pynput.keyboard import Controller as KeyboardController
import logging
from difflib import SequenceMatcher
from helpers import GrammarChecker
from models import DocumentContent, RetypedDocument
from typer import RealKeyboardTyper


from analyzer import analyze_docx

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class VerificationResult(BaseModel):
    """Class to store document verification results"""
    match_percentage: float
    differences: List[str] = []
    success: bool = False
    
    @property
    def is_successful(self):
        """Return True if match percentage is above 95%"""
        return self.match_percentage >= 95.0

class DocumentRetyper:
    """Class to retype documents with real keyboard typing"""
    def __init__(self):
        self.document_retyper = None
        self.keyboard_typer = RealKeyboardTyper(delay=0.01)  # Slightly slower for reliability
        self.keyboard = KeyboardController()
        self.last_typing_position = None
        self.last_typed_content = ""
        
    async def async_init(self):
        """Async initialization method"""
        self.document_retyper = Agent(
            # 'google-gla:gemini-2.0-flash',
            'groq:llama-3.3-70b-versatile',
            result_type=RetypedDocument,
            system_prompt=(
                "You are a document retyper assistant. Your task is to carefully read and retype"
                " documents, preserving their exact original format and content. "
                "BOLD AND ITALIC TEXT SHOULD BE PRESERVED. "
                 "For formatting, use the following markers:"
            " **bold text** for bold,"
            " __underlined text__ for underline,"
            " _italic text_ for italic."
            "use the addition functions like the ""addition"" function like toggle_bold, toggle_underline,toggle_italic "
                "USE THE REQUIRED TOOLS TO MAKE A CLEAR PROGRESS AND FASTER ERROR FREE RETYPING."
                "BULLETS, NUMBERING LIKE 1.2.3 ETC, AND PARAGRAPH SPACING SHOULD BE PRESERVED."
                "TYPE EXACT HOW THE DOCUMENT IS NO MORE STAFF ADDED IN IT AND BE CAREFUL."
                "Do not analyze, summarize, or modify the document in any way. Do not skip empty lines"
                " or spaces. Simply reproduce the document exactly as provided, maintaining all"
                "formatting, HEADINGS, bullet points, paragraph breaks, and structure."
                "Make sure that the Numbering and Bullet Points are preserved. And also incase ` is used in the document"
                " Make sure to preserve all newlines and paragraph spacing exactly as in the original."
            ),
        )
        return self

    async def verify_content(self, original_content: str, typed_content: str) -> VerificationResult:
        """Verify that typed content matches original content"""
        # Calculate similarity ratio
        similarity = SequenceMatcher(None, original_content, typed_content).ratio() * 100
        
        # Find differences if needed (simplified version)
        differences = []
        if similarity < 95.0:
            # Split by lines for easier comparison
            original_lines = original_content.split('\n')
            typed_lines = typed_content.split('\n')
            
            # Compare line by line
            for i, (orig_line, typed_line) in enumerate(zip(original_lines, typed_lines)):
                if orig_line != typed_line:
                    differences.append(f"Line {i+1}: Expected '{orig_line}', got '{typed_line}'")
            
            # If line counts differ
            if len(original_lines) != len(typed_lines):
                differences.append(f"Line count mismatch: Expected {len(original_lines)}, got {len(typed_lines)}")
        
        return VerificationResult(
            match_percentage=similarity,
            differences=differences[:10],  # Limit to first 10 differences
            success=similarity >= 95.0
        )
    
    async def get_current_content(self):
        """
        Method to get the currently typed content from the editor
        This is a placeholder - in a real implementation, you'd need 
        application-specific code to retrieve the content
        """
        # Simulate retrieving content - in reality you would:
        # 1. Use application APIs if available
        # 2. Or use keyboard shortcuts like Ctrl+A, Ctrl+C to copy content to clipboard
        # 3. Then read from clipboard
        
        # For demo purposes, prompt the user to provide the content
        print("\n>>> VERIFICATION STEP <<<")
        print("Please select all text in the editor (Ctrl+A), copy it (Ctrl+C)")
        input("Then press Enter to continue...")
        
        # In a real implementation, you would get clipboard content here
        import pyperclip  # You'll need to pip install pyperclip
        try:
            return pyperclip.paste()
        except Exception as e:
            print(f"Error getting clipboard content: {e}")
            return input("Please paste the document content manually: ")
    
    async def resume_typing_from_position(self, full_content: str, current_content: str):
        """Resume typing from where it left off based on current content"""
        # Find the position to resume from
        common_prefix_length = 0
        for i, (a, b) in enumerate(zip(current_content, full_content)):
            if a != b:
                break
            common_prefix_length = i + 1
        
        # Calculate what remains to be typed
        remaining_content = full_content[common_prefix_length:]
        
        if not remaining_content:
            print("Nothing left to type - content appears complete!")
            return
        
        print(f"Resuming typing from position {common_prefix_length}...")
        print(f"Approximately {len(remaining_content)} characters remaining to type.")
        
        # Resume typing the remaining content
        await self.keyboard_typer.type_with_verification(remaining_content, self.last_typing_position, True)
    
    async def recheck_and_correct_formatting(self, original_content: str, typed_content: str):
        """
        Identify and correct formatting issues between original and typed content
        """
        # Define formatting patterns
        bold_pattern = re.compile(r'\*\*(.*?)\*\*')
        italic_pattern = re.compile(r'_(.*?)_')
        underline_pattern = re.compile(r'__(.*?)__')
        
        # Check for missing formatting
        original_bolds = set(bold_pattern.findall(original_content))
        original_italics = set(italic_pattern.findall(original_content))
        original_underlines = set(underline_pattern.findall(original_content))
        
        typed_bolds = set(bold_pattern.findall(typed_content))
        typed_italics = set(italic_pattern.findall(typed_content))
        typed_underlines = set(underline_pattern.findall(typed_content))
        
        # Find missing formatting elements
        missing_bolds = original_bolds - typed_bolds
        missing_italics = original_italics - typed_italics
        missing_underlines = original_underlines - typed_underlines
        
        # Report issues
        format_issues = []
        if missing_bolds:
            format_issues.append(f"Missing bold formatting for: {', '.join(missing_bolds)}")
        if missing_italics:
            format_issues.append(f"Missing italic formatting for: {', '.join(missing_italics)}")
        if missing_underlines:
            format_issues.append(f"Missing underline formatting for: {', '.join(missing_underlines)}")
        
        if format_issues:
            print("Formatting issues detected:")
            for issue in format_issues:
                print(f"- {issue}")
            
            # Offer to correct formatting
            if input("Would you like to attempt to correct formatting issues? (y/n): ").lower() == 'y':
                await self.correct_formatting(typed_content, missing_bolds, missing_italics, missing_underlines)
        else:
            print("No formatting issues detected!")
            
    async def correct_formatting(self, current_text, missing_bolds, missing_italics, missing_underlines):
        """
        Helper method to apply missing formatting
        This is a simplified implementation that would need to be enhanced for real use
        """
        print("Formatting correction would happen here.")
        print("This would involve:")
        print("1. Finding each text segment that needs formatting")
        print("2. Placing cursor at the start of that segment")
        print("3. Applying the appropriate keyboard shortcuts")
        
        # In a real implementation, for each missing formatted item:
        # 1. Search for its position in the document
        # 2. Move cursor there
        # 3. Select the text
        # 4. Apply appropriate formatting shortcut (Ctrl+B for bold, etc.)
        
        for bold_text in missing_bolds:
            print(f"Would apply bold formatting to: '{bold_text}'")
            
        for italic_text in missing_italics:
            print(f"Would apply italic formatting to: '{italic_text}'")
            
        for underline_text in missing_underlines:
            print(f"Would apply underline formatting to: '{underline_text}'")
            
        print("Formatting correction complete.")

    async def retype_document_with_real_typing(self, document_text: str, typing_position=None):
        """Function to retype a document using real keyboard inputs with formatting support"""
        content = DocumentContent(text=document_text)
        
        # Store the typing position and original content for potential resumption
        self.last_typing_position = typing_position
        self.original_content = document_text
        
        # Just show a simple message in terminal
        print(f"Reading document content...")
        
        # Run the agent silently
        result = await self.document_retyper.run(
            f"Please retype the following document exactly as it is, preserving all formatting, spacing,"
            f" line breaks, and paragraph structure: {content.text}"
        )
        
        # Prepare user for typing
        print(f"Preparing to type the document with real keyboard inputs.")
        print(f"Please open the target file/editor and place your cursor where typing should begin.")
        print(f"You have 5 seconds to get ready...")
        
        # Countdown
        for i in range(5, 0, -1):
            print(f"{i}...")
            await asyncio.sleep(1)
        
        print("Typing now! You can move the mouse, the system will try to maintain focus...")
        
        # Perform real keyboard typing with formatting support
        try:
            # Check if the content likely has formatting markers
            if "**" in result.data.content or "__" in result.data.content or "_" in result.data.content:
                # Use the enhanced typing with formatting method
                await self.keyboard_typer.type_with_formatting(result.data.content, typing_position)
            else:
                # Use the regular typing with verification method
                await self.keyboard_typer.type_with_verification(result.data.content, typing_position, True)
            
            # Store what we've typed
            self.last_typed_content = result.data.content
            
            print(f"\nDocument typing complete! Now verifying...")
            
            # Perform verification
            await self.verify_retyped_document(document_text)
            
            print(f"\nDocument successfully retyped using real keyboard inputs!")
        except Exception as e:
            print(f"Error during typing: {str(e)}")
            traceback.print_exc()
            
            # Offer to resume typing
            if input("Would you like to try resuming typing from where it left off? (y/n): ").lower() == 'y':
                current_content = await self.get_current_content()
                await self.resume_typing_from_position(result.data.content, current_content)
        
        return result.data.content
    
    async def verify_retyped_document(self, original_content: str):
        """Function to verify retyped content against original"""
        print("\nVerifying document accuracy...")
        print("Please wait while we check the typed content...")
        
        # Get the current content from the editor
        typed_content = await self.get_current_content()
        
        # Verify content
        verification_result = await self.verify_content(original_content, typed_content)
        
        # Display results
        print(f"\nVerification Results:")
        print(f"Content Match: {verification_result.match_percentage:.2f}%")
        
        if verification_result.is_successful:
            print("✓ Document verification successful! The content matches the original.")
        else:
            print("⚠ Document verification found differences:")
            for diff in verification_result.differences:
                print(f"- {diff}")
            
            # Offer to correct differences
            if input("Would you like to attempt to correct differences? (y/n): ").lower() == 'y':
                await self.resume_typing_from_position(original_content, typed_content)
        
        # Check formatting specifically
        print("\nVerifying document formatting...")
        await self.recheck_and_correct_formatting(original_content, typed_content)

    async def display_document_info(self, file_path: str):
        """Display basic document info without typing the content"""
        # Check if it's a docx file

        file_ext = analyze_docx(file_path)
        print(f"File extension: {file_ext}")

        if file_path.endswith('.docx'):
            import docx
            doc = docx.Document(file_path)
            # Count paragraphs and characters in docx
            paragraphs = len(doc.paragraphs)
            char_count = sum(len(paragraph.text) for paragraph in doc.paragraphs)
            print(f"Document contains {paragraphs} paragraphs and {char_count} characters.")
            print("Starting retyping process...\n")
        else:
            # For regular text files
            with open(file_path, 'r') as f:
                content = f.read()
                
            print(f"Document loaded: {file_path}")
            line_count = content.count('\n') + 1
            char_count = len(content)
            print(f"Document contains {line_count} lines and {char_count} characters.")
            print("Starting retyping process...\n")

# Function to extract text from docx with better formatting preservation
def extract_text_from_docx(docx_path):
    """Extract text from docx file with enhanced formatting preservation"""
    import docx
    doc = docx.Document(docx_path)
    full_text = []
    
    # Process paragraphs with enhanced formatting
    for para in doc.paragraphs:
        # Don't skip empty paragraphs - they're part of formatting
        # Handle heading styles
        if para.style.name.startswith('Heading'):
            # Add appropriate formatting based on heading level
            level = int(para.style.name.replace('Heading', '')) if para.style.name != 'Heading' else 1
            prefix = '#' * level + ' ' if level > 0 else ''
            full_text.append(f"{prefix}{para.text}")
        else:
            # Regular paragraph - include even if empty to preserve spacing
            full_text.append(para.text)
    
    # Process tables if any
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            table_rows.append(" | ".join(row_text))
        full_text.append("\n".join(table_rows))
    
    # Join with paragraph spacing preserved
    return "\n\n".join(full_text)

# Example usage
if __name__ == "__main__":
    import asyncio
    import os
    import docx
    import random
    from docx import Document
    
    async def main():
        # Create and initialize the processor
        retyper = DocumentRetyper()
        await retyper.async_init()
        
        # Handle command line arguments if provided
        if len(sys.argv) > 1:
            file_path = sys.argv[1]
            if not os.path.exists(file_path):
                print(f"Error: File {file_path} not found.")
                return
            
            output_file = "retyped_output.txt"
            if len(sys.argv) > 2:
                output_file = sys.argv[2]
        else:
            # Default sample document path
            file_path = "doc.docx"
            output_file = "my_data.txt"
            
        # Show basic info about the document
        await retyper.display_document_info(file_path)
        
        # Get document text from file
        if file_path.endswith('.docx'):
            document_text = extract_text_from_docx(file_path)
        else:
            with open(file_path, 'r') as f:
                document_text = f.read()
        
        # Prepare for typing
        print("Please open a text editor where you want the typing to occur.")
        input("Press Enter when you're ready to begin typing...")
        
        # Get cursor position (optional)
        print("Place your cursor where typing should begin and keep it there.")
        print("Typing will start in 3 seconds...")
        await asyncio.sleep(3)
        
        # Current mouse position will be used as the focus point
        mouse = MouseController()
        cursor_position = mouse.position
        
        # Run the retyping directly to file
        try:
            retyped_content = await retyper.retype_document_with_real_typing(document_text, cursor_position)
            print(f"Document has been retyped with real keyboard inputs!")
        except Exception as e:
            print(f"Error occurred: {str(e)}")
            traceback.print_exc()
    
    asyncio.run(main())