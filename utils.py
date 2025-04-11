import mimetypes
import logging
from pathlib import Path
from typing import List
import chardet
import docx

import streamlit as st
from pynput.mouse import Button, Controller as MouseController
import io

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


#File verification
def is_valid_doc(file):
    """Validate if file is a proper document type with valid content"""
    try:
        file_extension = file.name.split('.')[-1].lower()
        mime_type = mimetypes.guess_type(file.name)[0]
        
        
        file_size = len(file.getvalue())
        if file_size > 10 * 1024 * 1024:  # 10MB
            return False, "File too large (max 10MB)"
        
        
        if file_extension == 'docx':
            try:
                doc = docx.Document(io.BytesIO(file.getvalue()))
                
                has_content = any(len(para.text.strip()) > 0 for para in doc.paragraphs)
                if not has_content:
                    return False, "DOCX file appears to be empty"
                return True, {"type": "docx", "mime": mime_type, "size": file_size}
            except Exception as e:
                return False, f"Invalid DOCX file: {str(e)}"
                
        elif file_extension in ['txt', 'py', 'md', 'pdf']:
            try:
                
                raw_data = file.getvalue()
                detected = chardet.detect(raw_data)
                encoding = detected['encoding'] if detected['encoding'] else 'utf-8'
                
                content = raw_data.decode(encoding)
                if not content.strip():
                    return False, "File appears to be empty"
                return True, {"type": file_extension, "mime": mime_type, "size": file_size, "encoding": encoding}
            except Exception as e:
                return False, f"Error reading file: {str(e)}"
        else:
            return False, f"Unsupported file type: {file_extension}"
            
    except Exception as e:
        logger.error(f"Validation error: {str(e)}")
        return False, f"Validation error: {str(e)}"


def chunk_content(content, chunk_size):
    """Split large content into manageable chunks"""
    chunks = []
    paragraphs = content.split("\n")
    current_chunk = ""
    
    for para in paragraphs:
        if len(current_chunk) + len(para) + 1 <= chunk_size:
            if current_chunk:
                current_chunk += "\n" + para
            else:
                current_chunk = para
        else:
            
            if len(para) > chunk_size:
                words = para.split(" ")
                temp_chunk = ""
                for word in words:
                    if len(temp_chunk) + len(word) + 1 <= chunk_size:
                        if temp_chunk:
                            temp_chunk += " " + word
                        else:
                            temp_chunk = word
                    else:
                        if current_chunk:
                            chunks.append(current_chunk)
                        current_chunk = temp_chunk
                        temp_chunk = word
                if temp_chunk:
                    current_chunk = temp_chunk
            else:
                
                chunks.append(current_chunk)
                current_chunk = para
    
   
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks


def display_screenshots(screenshots):
    """Display screenshots with proper formatting"""
    if not screenshots:
        st.warning("No screenshots captured during processing")
        return
    
    for i, screenshot in enumerate(screenshots):
        try:
            
            st.image(
                screenshot["data"],
                caption=screenshot.get("caption", f"Screenshot {i+1}"),
                use_column_width=True
            )
        except Exception as e:
            st.error(f"Error displaying screenshot {i+1}: {str(e)}")


# Extract text content from different file types
def extract_file_content(file, file_info):
    """Extract text content from various file types"""
    try:
        if file_info["type"] == "docx":
            doc = docx.Document(io.BytesIO(file.getvalue()))
            return "\n".join([para.text for para in doc.paragraphs])
        else:  # txt, py, md
            encoding = file_info.get("encoding", "utf-8")
            try:
                return file.getvalue().decode(encoding)
            except UnicodeDecodeError:
                # Fallback to latin-1 if detected encoding fails
                return file.getvalue().decode('latin-1', errors='replace')
    except Exception as e:
        logger.error(f"Content extraction error: {str(e)}")
        raise Exception(f"Failed to extract content: {str(e)}")
