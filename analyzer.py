from docx import Document
from docx2python import docx2python
import os
import io

def analyze_docx(file_path) -> str:
    analysis_result = []
    
    doc = Document(file_path)
    
    analysis_result.append("== PARAGRAPHS AND STYLES ==")
    for para in doc.paragraphs:
        if para.text.strip():
            analysis_result.append(f"Text: {para.text}")
            analysis_result.append(f"Style: {para.style.name}")
            analysis_result.append("-" * 50)
    
    analysis_result.append("\n== TABLES ==")
    for table_idx, table in enumerate(doc.tables):
        analysis_result.append(f"Table {table_idx + 1}:")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            analysis_result.append("\t".join(row_data))
        analysis_result.append("-" * 50)
    
    try:
        core_props = doc.core_properties
        analysis_result.append("\n== METADATA ==")
        analysis_result.append(f"Author: {core_props.author}")
        analysis_result.append(f"Created: {core_props.created}")
        analysis_result.append(f"Modified: {core_props.modified}")
        analysis_result.append("-" * 50)
    except Exception as e:
        analysis_result.append(f"\n== METADATA EXTRACTION ERROR: {str(e)} ==")
    
    try:
        doc_content = docx2python(file_path)
        
        analysis_result.append("\n== HEADERS ==")
        for header in doc_content.header:
            if header:
                analysis_result.append(str(header))
        analysis_result.append("-" * 50)
        
        analysis_result.append("\n== FOOTERS ==")
        for footer in doc_content.footer:
            if footer:
                analysis_result.append(str(footer))
        analysis_result.append("-" * 50)
    except Exception as e:
        analysis_result.append(f"\n== HEADERS/FOOTERS EXTRACTION ERROR: {str(e)} ==")
    
    try:
        analysis_result.append("\n== IMAGES ==")
        image_count = 0
        
        
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image_part = doc.part.rels[rel].target_part
                image_filename = os.path.basename(image_part.partname)
                image_count += 1
                analysis_result.append(f"Image {image_count}: {image_filename}")
        
        if image_count == 0:
            analysis_result.append("No images found in document")
        analysis_result.append("-" * 50)
    except Exception as e:
        analysis_result.append(f"\n== IMAGE EXTRACTION ERROR: {str(e)} ==")
    
    
    return "\n".join(analysis_result)


"""
with st.expander("🔑 Analyzer Formater", expanded=False):
    file_analyzed = analyze_docx(uploaded_file)
    st.markdown("#### 🧾 File Content:")

    if validation_result['type'] == 'py':
        # Use code highlighting for code files
        st.code(file_analyzed, language=validation_result['type'])
    else:
        # Use scrollable text area for documents
        st.text_area(
            label="Document Contains the following Formats as Word Document Provided:",
            value=file_analyzed,
            height=400,
            disabled=True
        )
st.divider()
"""