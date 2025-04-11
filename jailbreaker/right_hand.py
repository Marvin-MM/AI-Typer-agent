import os
import time
import sys
import asyncio
from pydantic_ai import Agent
#Groq 
from pydantic_ai.models.groq import GroqModel
from dotenv import load_dotenv
import traceback
from pynput.mouse import Button, Controller as MouseController
import logging
from models import DocumentContent, RetypedDocument
from typer import RealKeyboardTyper


from analyzer import analyze_docx

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentRetyper:
    """Class to retype documents with real keyboard typing"""
    def __init__(self):
        self.document_retyper = None
        self.keyboard_typer = RealKeyboardTyper(delay=0.03, verify_interval=50)
        self.focus_position = (500,500)  # Slightly slower for reliability
        
    async def async_init(self):
        """Async initialization method"""
        self.document_retyper = Agent(
            # 'google-gla:gemini-2.0-flash',
            'groq:llama-3.3-70b-versatile',
            result_type=RetypedDocument,
            system_prompt=(
                "You are a document retyper assistant. Your task is to carefully read and retype"
                " documents, preserving their exact original format and content. "
                "BOLD AND ITALIC TEXT SHOULD BE PRESERVED. "
                 "For formatting, use the following markers:"
            " **bold text** for bold,"
            " __underlined text__ for underline,"
            " _italic text_ for italic."
            "use the addition functions like the ""addition"" function like toggle_bold, toggle_underline,toggle_italic "
                "USE THE REQUIRED TOOLS TO MAKE A CLEAR PROGRESS AND FASTER ERROR FREE RETYPING."
                "BULLETS, NUMBERING LIKE 1.2.3 ETC, AND PARAGRAPH SPACING SHOULD BE PRESERVED."
                "TYPE EXACT HOW THE DOCUMENT IS NO MORE STAFF ADDED IN IT AND BE CAREFUL."
                "Do not analyze, summarize, or modify the document in any way. Do not skip empty lines"
                " or spaces. Simply reproduce the document exactly as provided, maintaining all"
                "formatting, HEADINGS, bullet points, paragraph breaks, and structure."
                "Make sure that the Numbering and Bullet Points are preserved. And also incase ` is used in the document"
                " Make sure to preserve all newlines and paragraph spacing exactly as in the original."
            ),
        )
        return self

    async def retype_document_with_real_typing(self, document_text: str, typing_position=None):
        """Function to retype a document using real keyboard inputs with formatting support"""
        content = DocumentContent(text=document_text)
        
        # Just show a simple message in terminal
        print(f"Reading document content...")
        
        # Run the agent silently
        result = await self.document_retyper.run(
            f"Please retype the following document exactly as it is, preserving all formatting, spacing,"
            f" line breaks, and paragraph structure: {content.text}"
        )
        
        # Prepare user for typing
        print(f"Preparing to type the document with real keyboard inputs.")
        print(f"Please open the target file/editor and place your cursor where typing should begin.")
        print(f"You have 5 seconds to get ready...")
        
        # Countdown
        for i in range(5, 0, -1):
            print(f"{i}...")
            await asyncio.sleep(1)
        
        print("Typing now! You can move the mouse, the system will try to maintain focus...")
        
        # Perform real keyboard typing with formatting support
        try:
            # Check if the content likely has formatting markers
            if "**" in result.data.content or "__" in result.data.content or "_" in result.data.content:
                # Use the enhanced typing with formatting method
                await self.keyboard_typer.type_with_formatting(result.data.content, typing_position)
            else:
                # Use the regular typing with verification method
                await self.keyboard_typer.type_with_verification(result.data.content, typing_position,  True)
            print(f"\nDocument successfully retyped using real keyboard inputs!")
        except Exception as e:
            print(f"Error during typing: {str(e)}")
            traceback.print_exc()
        
        return result.data.content

    async def display_document_info(self, file_path: str):
        """Display basic document info without typing the content"""
        # Check if it's a docx file

        file_ext = analyze_docx(file_path)
        print(f"File extension: {file_ext}")

        if file_path.endswith('.docx'):
            import docx
            doc = docx.Document(file_path)
            # Count paragraphs and characters in docx
            paragraphs = len(doc.paragraphs)
            char_count = sum(len(paragraph.text) for paragraph in doc.paragraphs)
            print(f"Document contains {paragraphs} paragraphs and {char_count} characters.")
            print("Starting retyping process...\n")
        else:
            # For regular text files
            with open(file_path, 'r') as f:
                content = f.read()
                
            print(f"Document loaded: {file_path}")
            line_count = content.count('\n') + 1
            char_count = len(content)
            print(f"Document contains {line_count} lines and {char_count} characters.")
            print("Starting retyping process...\n")

# Function to extract text from docx with better formatting preservation
def extract_text_from_docx(docx_path):
    """Extract text from docx file with enhanced formatting preservation"""
    import docx
    doc = docx.Document(docx_path)
    full_text = []
    
    # Process paragraphs with enhanced formatting
    for para in doc.paragraphs:
        # Don't skip empty paragraphs - they're part of formatting
        # Handle heading styles
        if para.style.name.startswith('Heading'):
            # Add appropriate formatting based on heading level
            level = int(para.style.name.replace('Heading', '')) if para.style.name != 'Heading' else 1
            prefix = '#' * level + ' ' if level > 0 else ''
            full_text.append(f"{prefix}{para.text}")
        else:
            # Regular paragraph - include even if empty to preserve spacing
            full_text.append(para.text)
    
    # Process tables if any
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            table_rows.append(" | ".join(row_text))
        full_text.append("\n".join(table_rows))
    
    # Join with paragraph spacing preserved
    return "\n\n".join(full_text)

# Example usage
if __name__ == "__main__":
    import asyncio
    import os
    import docx
    import random
    from docx import Document
    
    async def main():
        # Create and initialize the processor
        retyper = DocumentRetyper()
        await retyper.async_init()
        
        # Handle command line arguments if provided
        if len(sys.argv) > 1:
            file_path = sys.argv[1]
            if not os.path.exists(file_path):
                print(f"Error: File {file_path} not found.")
                return
            
            output_file = "retyped_output.txt"
            if len(sys.argv) > 2:
                output_file = sys.argv[2]
        else:
            # Default sample document path
            file_path = "doc.docx"
            output_file = "my_data.txt"
            
        # Show basic info about the document
        await retyper.display_document_info(file_path)
        
        # Get document text from file
        if file_path.endswith('.docx'):
            document_text = extract_text_from_docx(file_path)
        else:
            with open(file_path, 'r') as f:
                document_text = f.read()
        
        # Prepare for typing
        print("Please open a text editor where you want the typing to occur.")
        input("Press Enter when you're ready to begin typing...")
        
        # Get cursor position (optional)
        print("Place your cursor where typing should begin and keep it there.")
        print("Typing will start in 3 seconds...")
        await asyncio.sleep(3)
        
        # Current mouse position will be used as the focus point
        mouse = MouseController()
        cursor_position = mouse.position
        
        # Run the retyping directly to file
        try:
            retyped_content = await retyper.retype_document_with_real_typing(document_text, cursor_position)
            print(f"Document has been retyped with real keyboard inputs!")
        except Exception as e:
            print(f"Error occurred: {str(e)}")
            traceback.print_exc()
    
    asyncio.run(main())