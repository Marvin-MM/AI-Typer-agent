
import os
import time
import sys
import asyncio
from typing import List, Optional, Dict, Any
import re
from pydantic import BaseModel, Field
from pydantic_ai import Agent
from dotenv import load_dotenv
import traceback
import random
from pynput.keyboard import Key, Controller as KeyboardController
from pynput.mouse import Button, Controller as MouseController
import keyboard
import language_tool_python

load_dotenv()

# Initialize language tool for grammar checking
try:
    language_tool = language_tool_python.LanguageTool('en-US')
except:
    print("Warning: LanguageTool could not be initialized. Grammar checking will be disabled.")
    language_tool = None


class GrammarChecker:
    """Class to check grammar and provide corrections"""
    
    def __init__(self):
        self.language_tool = language_tool
    
    def check_grammar(self, text: str) -> List[Dict]:
        """Check grammar and return corrections"""
        if not self.language_tool:
            return []
        
        # Get matches from language tool
        matches = self.language_tool.check(text)
        
        corrections = []
        for match in matches:
            if match.replacements:
                corrections.append({
                    "type": "grammar",
                    "message": match.message,
                    "offset": match.offset,
                    "length": match.errorLength,
                    "replacements": match.replacements,
                    "rule_id": match.ruleId
                })
        
        return corrections
    
    def apply_correction(self, text: str, correction: Dict) -> str:
        """Apply a single correction to the text"""
        if correction["type"] == "grammar" and correction["replacements"]:
            # Replace the text at the given offset with the first suggestion
            offset = correction["offset"]
            length = correction["length"]
            replacement = correction["replacements"][0]
            
            return text[:offset] + replacement + text[offset + length:]
        
        return text


class DocumentContent(BaseModel):
    """Structure for the document to be retyped"""
    text: str = Field(description="Raw text of the document")
    source: Optional[str] = Field(None, description="Source of the document")

class RetypedDocument(BaseModel):
    """Structure for the retyped document"""
    content: str = Field(description="The exact content of the document, retaining original formatting")

class DocumentVerifier:
    """Class to verify document content and detect errors"""
    def __init__(self):
        pass
    
    def compare_content(self, original: str, typed: str) -> Dict[str, Any]:
        """Compare original content with typed content to detect errors"""
        # Split into lines for comparison
        original_lines = original.split('\n')
        typed_lines = typed.split('\n')
        
        errors = []
        
        # Check length difference first
        if len(original_lines) != len(typed_lines):
            errors.append({
                "type": "line_count_mismatch",
                "message": f"Line count mismatch: {len(original_lines)} vs {len(typed_lines)}"
            })
        
        # Compare line by line
        for i, (orig, typed) in enumerate(zip(original_lines, typed_lines[:len(original_lines)])):
            if orig != typed:
                # Find the position of the first difference
                pos = next((j for j in range(min(len(orig), len(typed))) if orig[j] != typed[j]), min(len(orig), len(typed)))
                errors.append({
                    "type": "content_mismatch",
                    "line": i + 1,
                    "position": pos + 1,
                    "original": orig,
                    "typed": typed
                })
        
        return {
            "match": len(errors) == 0,
            "errors": errors
        }

class RealKeyboardTyper:
    """Class to perform real keyboard typing using pynput with error correction"""
    def __init__(self, delay=0.01, verify_interval=100):  # Slightly slower for reliability
        self.keyboard = KeyboardController()
        self.mouse = MouseController()
        self.grammar_checker = GrammarChecker()
        self.delay = delay
        self.verify_interval = verify_interval  # Check every X characters
        self.document_verifier = DocumentVerifier()

    def toggle_bold(self):
        """Toggle bold formatting using Ctrl+B shortcut"""
        with self.keyboard.pressed(Key.ctrl):
            self.keyboard.press('b')
            self.keyboard.release('b')
        time.sleep(0.1)

    def toggle_underline(self):
        """Toggle underline formatting using Ctrl+U shortcut"""
        with self.keyboard.pressed(Key.ctrl):
            self.keyboard.press('u')
            self.keyboard.release('u')
        time.sleep(0.1)

    def toggle_italic(self):
        """Toggle italic formatting using Ctrl+I shortcut"""
        with self.keyboard.pressed(Key.ctrl):
            self.keyboard.press('i')
            self.keyboard.release('i')
        time.sleep(0.1)

    def click_and_focus(self, position=None):
        """Click at specified coordinates to focus or use current position"""
        if position:
            # Save current position
            old_pos = self.mouse.position
            # Move to target position
            self.mouse.position = position
            time.sleep(0.1)
            # Click to focus
            self.mouse.click(Button.left)
            time.sleep(0.1)
            # Return to original position to avoid interference
            self.mouse.position = old_pos
        else:
            # Just click at current position
            self.mouse.click(Button.left)
            time.sleep(0.1)
    
    def select_all_and_delete(self):
        """Select all text and delete it to start fresh"""
        with self.keyboard.pressed(Key.ctrl):
            self.keyboard.press('a')
            self.keyboard.release('a')
        time.sleep(0.1)
        self.keyboard.press(Key.delete)
        self.keyboard.release(Key.delete)
        time.sleep(0.1)

    async def type_with_formatting(self, text, focus_position=None):
        """Type text with formatting support for bold, italic, and underline"""
        # Initial focus
        self.click_and_focus(focus_position)
        
        # Process text with formatting markers
        # For example, you could use markers like **bold**, _italic_, __underline__
        
        i = 0
        while i < len(text):
            # Check for bold markers
            if text[i:i+2] == "**" and "**" in text[i+2:]:
                # Skip the opening marker
                i += 2
                # Find the closing marker
                end_bold = text.find("**", i)
                if end_bold != -1:
                    # Get the text to be bolded
                    bold_text = text[i:end_bold]
                    # Toggle bold on
                    self.toggle_bold()
                    # Type the text
                    for char in bold_text:
                        if not self.handle_special_char(char):
                            self.keyboard.press(char)
                            self.keyboard.release(char)
                        await asyncio.sleep(self.delay)
                    # Toggle bold off
                    self.toggle_bold()
                    # Skip to after the closing marker
                    i = end_bold + 2
                    continue
            
            # Check for underline markers (similar logic)
            if text[i:i+2] == "__" and "__" in text[i+2:]:
                # Skip the opening marker
                i += 2
                # Find the closing marker
                end_underline = text.find("__", i)
                if end_underline != -1:
                    # Get the text to be underlined
                    underline_text = text[i:end_underline]
                    # Toggle underline on
                    self.toggle_underline()
                    # Type the text
                    for char in underline_text:
                        if not self.handle_special_char(char):
                            self.keyboard.press(char)
                            self.keyboard.release(char)
                        await asyncio.sleep(self.delay)
                    # Toggle underline off
                    self.toggle_underline()
                    # Skip to after the closing marker
                    i = end_underline + 2
                    continue
            
            # Type regular character
            if not self.handle_special_char(text[i]):
                self.keyboard.press(text[i])
                self.keyboard.release(text[i])
            
            # Move to next character
            i += 1
            await asyncio.sleep(self.delay)
    
    def handle_special_char(self, char):
        """Handle typing of special characters"""
        if char == '\n':
            self.keyboard.press(Key.enter)
            self.keyboard.release(Key.enter)
            return True
        elif char == '\t':
            self.keyboard.press(Key.tab)
            self.keyboard.release(Key.tab)
            return True
        elif char == ' ':
            self.keyboard.press(Key.space)
            self.keyboard.release(Key.space)
            return True
        return False
    
    def handle_text_selection(self, count=1, direction="forward"):
        """Handle text selection for correction"""
        key = Key.right if direction == "forward" else Key.left
        with self.keyboard.pressed(Key.shift):
            for _ in range(count):
                self.keyboard.press(key)
                self.keyboard.release(key)
                time.sleep(0.01)
    
    def preserve_formatting_boundaries(self, text):
        """Ensure proper formatting boundaries are preserved"""
        # Make sure paragraph breaks are preserved
        preserved = text
        # Replace single newlines between paragraphs with double newlines if not already
        preserved = re.sub(r'([^\n])\n([^\n])', r'\1\n\n\2', preserved)
        return preserved
    
    async def apply_corrections(self, expected, actual):
        """Apply corrections to make actual text match expected text"""
        if expected == actual:
            return
        
        # Calculate differences
        min_len = min(len(expected), len(actual))
        
        # Find the first position where they differ
        diff_pos = 0
        for i in range(min_len):
            if expected[i] != actual[i]:
                diff_pos = i
                break
        else:
            # If no difference found in common length, the difference is in length
            diff_pos = min_len
        
        # Calculate how many characters to delete and what to type
        chars_to_delete = len(actual) - diff_pos
        chars_to_type = expected[diff_pos:]
        
        if chars_to_delete > 0:
            # Select and delete the incorrect characters
            for _ in range(chars_to_delete):
                self.keyboard.press(Key.backspace)
                self.keyboard.release(Key.backspace)
                await asyncio.sleep(0.01)
        
        # Type the correct characters
        for char in chars_to_type:
            if not self.handle_special_char(char):
                self.keyboard.press(char)
                self.keyboard.release(char)
            await asyncio.sleep(self.delay)
    async def type_with_verification(self, text, focus_position=None, error_correction=True):
        """Type text with periodic verification and error correction"""
        # Initial focus
        self.click_and_focus(focus_position)
        
        # Ensure proper formatting boundaries
        text = self.preserve_formatting_boundaries(text)
        
        # Variables for cursor tracking
        cursor_line = 0
        cursor_pos = 0
        typed_content = ""
        
        # Type character by character with verification
        char_count = 0
        segments = text.split('\n')
        
        for i, segment in enumerate(segments):
            # Type each character in the segment
            for j, char in enumerate(segment):
                # Update cursor position tracking
                cursor_pos = j
                
                # Regular typing with error handling
                try:
                    # Type the character
                    self.keyboard.press(char)
                    self.keyboard.release(char)
                    typed_content += char
                except Exception as e:
                    print(f"Error typing character '{char}': {str(e)}")
                    # Try alternative method for problematic characters
                    try:
                        keyboard.write(char)
                        typed_content += char
                    except:
                        print(f"Failed to type character '{char}'")
                
                # Dynamic delay between keystrokes for realistic typing
                typing_delay = self.delay + (0.01 * (0.5 - random.random()))
                await asyncio.sleep(typing_delay)
                
                # Verify every X characters if enabled
                char_count += 1
                if error_correction and char_count % self.verify_interval == 0:
                    await self.verify_and_correct(text[:len(typed_content)], typed_content, cursor_line, cursor_pos)

                if error_correction and char_count % self.verify_interval == 0:
                            expected_so_far = text[:len(typed_content)]
                            await self.apply_corrections(expected_so_far, typed_content)
                            
                            # Check grammar in the last sentence
                            last_sentence = re.split(r'[.!?]', typed_content)[-1]
                            if len(last_sentence) > 10:  # Only check substantial sentences
                                corrections = self.grammar_checker.check_grammar(last_sentence)
                                if corrections:
                                    print(f"Grammar correction suggested: {corrections[0]['message']}")
                                    # Grammar corrections disabled by default as they change content
                                    # Uncomment to enable automatic grammar fixes
                        
            # Add newline if not the last segment
            if i < len(segments) - 1:
                self.keyboard.press(Key.enter)
                self.keyboard.release(Key.enter)
                typed_content += '\n'
                cursor_line += 1
                cursor_pos = 0
                await asyncio.sleep(self.delay)
        
        # Final verification
        if error_correction:
            await self.verify_and_correct(text, typed_content, cursor_line, cursor_pos)
        
        return typed_content
    
    def preserve_formatting_boundaries(self, text):
        """Ensure proper formatting boundaries are preserved"""
        # Make sure paragraph breaks are preserved
        preserved = text
        # Replace single newlines between paragraphs with double newlines if not already
        preserved = re.sub(r'([^\n])\n([^\n])', r'\1\n\n\2', preserved)
        return preserved
    
    
    async def verify_and_correct(self, expected, actual, current_line, current_pos):
        """Verify typed content and attempt to correct errors"""
        # Save original cursor position
        original_line = current_line
        original_pos = current_pos
        
        # Skip verification if strings match exactly
        if expected == actual:
            return
        
        # Simple check for line count discrepancy
        expected_lines = expected.count('\n') + 1
        actual_lines = actual.count('\n') + 1
        
        if expected_lines != actual_lines:
            print(f"Warning: Line count mismatch - expected {expected_lines}, got {actual_lines}")
            
            # Try to correct by moving to the end and adding needed newlines
            for _ in range(abs(expected_lines - actual_lines)):
                self.keyboard.press(Key.enter)
                self.keyboard.release(Key.enter)
                await asyncio.sleep(0.05)
        
        # Look for character mismatches and try to correct them
        comparison = self.document_verifier.compare_content(expected, actual)
        
        if not comparison["match"] and len(comparison["errors"]) > 0:
            print(f"Detected {len(comparison['errors'])} errors, attempting to correct...")
            
            # We'll implement a simple correction for the first error
            if len(comparison["errors"]) > 0:
                err = comparison["errors"][0]
                if err["type"] == "content_mismatch":
                    # Try to navigate to the error position
                    # This is a simplified approach - more sophisticated navigation would be needed
                    # for real-world applications
                    
                    # First press Escape to ensure we're not in any special mode
                    self.keyboard.press(Key.esc)
                    self.keyboard.release(Key.esc)
                    await asyncio.sleep(0.1)
                    
                    # Try to correct by retyping the problematic part
                    print(f"Attempting recovery...")
                    self.click_and_focus()
                    await asyncio.sleep(0.1)
                    
                    # Select all and delete as a recovery method
                    self.select_all_and_delete()
                    
                    # Retype from the beginning of the problematic section
                    # For simplicity, we'll just retype a portion of the text
                    recovery_text = expected[:min(len(expected), 200)]  # Just first 200 chars for demo
                    for recovery_char in recovery_text:
                        if not self.handle_special_char(recovery_char):
                            self.keyboard.press(recovery_char)
                            self.keyboard.release(recovery_char)
                        await asyncio.sleep(self.delay)

    async def type_text(self, text, focus_position=None):
        """Type the text using real keyboard inputs with proper focus management"""
        # Initial focus
        if focus_position:
            self.click_and_focus(focus_position)
        else:
            self.click_and_focus()
        
        # Set up a mouse position check interval to maintain focus
        last_check = time.time()
        check_interval = 1.0  # Check focus more frequently (every 1 second)
        
        # Remember initial position to check if mouse moved
        initial_pos = self.mouse.position
        
        # Keep track of what we've typed for verification
        typed_text = ""
        
        # Split text by paragraphs to handle formatting better
        paragraphs = re.split(r'\n\s*\n', text)
        
        for i, paragraph in enumerate(paragraphs):
            # Check if we need to refocus
            current_time = time.time()
            if current_time - last_check > check_interval:
                # Check if mouse has moved significantly
                current_pos = self.mouse.position
                distance = ((current_pos[0] - initial_pos[0])**2 + 
                           (current_pos[1] - initial_pos[1])**2)**0.5
                
                if distance > 5:  # Mouse moved more than 5 pixels
                    # Refocus by clicking
                    print("Mouse moved - refocusing...")
                    self.click_and_focus(initial_pos)
                    initial_pos = self.mouse.position  # Update initial position
                
                last_check = current_time
            
            # Type each character in the paragraph
            for char in paragraph:
                # Handle special characters
                if not self.handle_special_char(char):
                    # Type the character
                    self.keyboard.press(char)
                    self.keyboard.release(char)
                
                typed_text += char
                
                # Dynamic delay between keystrokes for realistic typing
                typing_delay = self.delay + (0.01 * (0.5 - random.random()))
                await asyncio.sleep(typing_delay)
            
            # Add paragraph breaks between paragraphs, except after the last one
            if i < len(paragraphs) - 1:
                # Add two newlines for paragraph breaks
                self.keyboard.press(Key.enter)
                self.keyboard.release(Key.enter)
                await asyncio.sleep(self.delay)
                self.keyboard.press(Key.enter)
                self.keyboard.release(Key.enter)
                await asyncio.sleep(self.delay)
                typed_text += "\n\n"
        
        return typed_text

class DocumentRetyper:
    """Class to retype documents with real keyboard typing"""
    def __init__(self):
        self.document_retyper = None
        self.keyboard_typer = RealKeyboardTyper(delay=0.03)  # Slightly slower for reliability
        
    async def async_init(self):
        """Async initialization method"""
        self.document_retyper = Agent(
            'google-gla:gemini-2.0-flash',
            result_type=RetypedDocument,
            system_prompt=(
                "You are a document retyper assistant. Your task is to carefully read and retype"
                " documents, preserving their exact original format and content. "
                "BOLD AND ITALIC TEXT SHOULD BE PRESERVED. "
                 "For formatting, use the following markers:"
            " **bold text** for bold,"
            " __underlined text__ for underline,"
            " _italic text_ for italic."
            "use the addition functions like the ""addition"" function like toggle_bold, toggle_underline,toggle_italic "
                "USE THE REQUIRED TOOLS TO MAKE A CLEAR PROGRESS AND FASTER ERROR FREE RETYPING."
                "BULLETS, NUMBERING LIKE 1.2.3 ETC, AND PARAGRAPH SPACING SHOULD BE PRESERVED."
                "TYPE EXACT HOW THE DOCUMENT IS NO MORE STAFF ADDED IN IT AND BE CAREFUL."
                "Do not analyze, summarize, or modify the document in any way. Do not skip empty lines"
                " or spaces. Simply reproduce the document exactly as provided, maintaining all"
                "formatting, HEADINGS, bullet points, paragraph breaks, and structure."
                "Make sure that the Numbering and Bullet Points are preserved. And also incase ` is used in the document"
                " Make sure to preserve all newlines and paragraph spacing exactly as in the original."
            ),
        )
        return self

    async def retype_document_with_real_typing(self, document_text: str, typing_position=None):
        """Function to retype a document using real keyboard inputs with formatting support"""
        content = DocumentContent(text=document_text)
        
        # Just show a simple message in terminal
        print(f"Reading document content...")
        
        # Run the agent silently
        result = await self.document_retyper.run(
            f"Please retype the following document exactly as it is, preserving all formatting, spacing,"
            f" line breaks, and paragraph structure: {content.text}"
        )
        
        # Prepare user for typing
        print(f"Preparing to type the document with real keyboard inputs.")
        print(f"Please open the target file/editor and place your cursor where typing should begin.")
        print(f"You have 5 seconds to get ready...")
        
        # Countdown
        for i in range(5, 0, -1):
            print(f"{i}...")
            await asyncio.sleep(1)
        
        print("Typing now! You can move the mouse, the system will try to maintain focus...")
        
        # Perform real keyboard typing with formatting support
        try:
            # Check if the content likely has formatting markers
            if "**" in result.data.content or "__" in result.data.content or "_" in result.data.content:
                # Use the enhanced typing with formatting method
                await self.keyboard_typer.type_with_formatting(result.data.content, typing_position)
            else:
                # Use the regular typing with verification method
                await self.keyboard_typer.type_with_verification(result.data.content, typing_position, True)
            print(f"\nDocument successfully retyped using real keyboard inputs!")
        except Exception as e:
            print(f"Error during typing: {str(e)}")
            traceback.print_exc()
        
        return result.data.content

    async def display_document_info(self, file_path: str):
        """Display basic document info without typing the content"""
        # Check if it's a docx file
        if file_path.endswith('.docx'):
            import docx
            doc = docx.Document(file_path)
            # Count paragraphs and characters in docx
            paragraphs = len(doc.paragraphs)
            char_count = sum(len(paragraph.text) for paragraph in doc.paragraphs)
            

            print(f"Document loaded: {file_path}")
            print(f"Document contains {paragraphs} paragraphs and {char_count} characters.")
            print("Starting retyping process...\n")
        else:
            # For regular text files
            with open(file_path, 'r') as f:
                content = f.read()
                
            print(f"Document loaded: {file_path}")
            line_count = content.count('\n') + 1
            char_count = len(content)
            print(f"Document contains {line_count} lines and {char_count} characters.")
            print("Starting retyping process...\n")

# Function to extract text from docx with better formatting preservation
def extract_text_from_docx(docx_path):
    """Extract text from docx file with enhanced formatting preservation"""
    import docx
    doc = docx.Document(docx_path)
    full_text = []
    
    # Process paragraphs with enhanced formatting
    for para in doc.paragraphs:
        # Don't skip empty paragraphs - they're part of formatting
        # Handle heading styles
        if para.style.name.startswith('Heading'):
            # Add appropriate formatting based on heading level
            level = int(para.style.name.replace('Heading', '')) if para.style.name != 'Heading' else 1
            prefix = '#' * level + ' ' if level > 0 else ''
            full_text.append(f"{prefix}{para.text}")
        else:
            # Regular paragraph - include even if empty to preserve spacing
            full_text.append(para.text)
    
    # Process tables if any
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            table_rows.append(" | ".join(row_text))
        full_text.append("\n".join(table_rows))
    
    # Join with paragraph spacing preserved
    return "\n\n".join(full_text)

# Example usage
if __name__ == "__main__":
    import asyncio
    import os
    import docx
    import random
    from docx import Document
    
    async def main():
        # Create and initialize the processor
        retyper = DocumentRetyper()
        await retyper.async_init()
        
        # Handle command line arguments if provided
        if len(sys.argv) > 1:
            file_path = sys.argv[1]
            if not os.path.exists(file_path):
                print(f"Error: File {file_path} not found.")
                return
            
            output_file = "retyped_output.txt"
            if len(sys.argv) > 2:
                output_file = sys.argv[2]
        else:
            # Default sample document path
            file_path = "doc.docx"
            output_file = "my_data.txt"
            
            # Check if sample document exists, create a simple one if not
            if not os.path.exists(file_path):
                print(f"Sample document not found at {file_path}, creating a simple one...")
                # Create a sample .docx file
                doc = Document()
                doc.add_heading('Sample Document', 0)
                doc.add_paragraph('This is a sample document for testing the document retyper.')
                
                # Add an empty paragraph for spacing
                doc.add_paragraph('')
                
                doc.add_heading('Introduction', level=2)
                doc.add_paragraph('This document serves as a test case for our document retyping tool.')
                
                # Add multiple paragraphs with different spacing
                doc.add_paragraph('This paragraph demonstrates the preservation of paragraph spacing.')
                doc.add_paragraph('This second paragraph should maintain proper distance from the previous one.')
                
                doc.add_heading('Key Features', level=2)
                
                # Add a bulleted list
                features = doc.add_paragraph()
                features.add_run('- Exact reproduction of text\n')
                features.add_run('- Preservation of formatting\n')
                features.add_run('- Maintenance of structure\n')
                features.add_run('- No analysis or modification')
                
                # Add another empty paragraph
                doc.add_paragraph('')
                
                doc.add_heading('Conclusion', level=2)
                doc.add_paragraph('This sample demonstrates how our tool can retype document content exactly as provided.')
                
                # Save the document
                doc.save(file_path)
        
        # Show basic info about the document
        await retyper.display_document_info(file_path)
        
        # Get document text from file
        if file_path.endswith('.docx'):
            document_text = extract_text_from_docx(file_path)
        else:
            with open(file_path, 'r') as f:
                document_text = f.read()
        
        # Prepare for typing
        print("Please open a text editor where you want the typing to occur.")
        input("Press Enter when you're ready to begin typing...")
        
        # Get cursor position (optional)
        print("Place your cursor where typing should begin and keep it there.")
        print("Typing will start in 3 seconds...")
        await asyncio.sleep(3)
        
        # Current mouse position will be used as the focus point
        mouse = MouseController()
        cursor_position = mouse.position
        
        # Run the retyping directly to file
        try:
            retyped_content = await retyper.retype_document_with_real_typing(document_text, cursor_position)
            print(f"Document has been retyped with real keyboard inputs!")
        except Exception as e:
            print(f"Error occurred: {str(e)}")
            traceback.print_exc()
    
    asyncio.run(main())